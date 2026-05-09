import docx

doc = docx.Document("Intenship001.docx")
print(f"Number of tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables[:2]):
    print(f"Table {i}:")
    for row in table.rows[:3]:
        print([cell.text.strip() for cell in row.cells])
