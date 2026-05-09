import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def copy_paragraph_style(target_p, source_style_name, doc_template):
    try:
        target_p.style = doc_template.styles[source_style_name]
    except:
        pass # fallback if style doesn't exist

def format_report():
    template_path = "Intenship001.docx"
    source_path = "Landslide report.docx"
    output_path = "Landslide_Report_Formatted.docx"

    if not os.path.exists(template_path) or not os.path.exists(source_path):
        print("Files missing")
        return

    template_doc = docx.Document(template_path)
    source_doc = docx.Document(source_path)
    
    # Create output doc by copying template to keep styles and headers
    output_doc = docx.Document(template_path)
    
    # Clear all paragraphs and tables from the copy (keeping headers/footers)
    # We leave the first paragraph because we can't delete all paragraphs
    for p in output_doc.paragraphs:
        p.text = ""
    
    # Actually, a better way is to clear body but keep sections
    # But python-docx doesn't make it easy to delete everything.
    # Let's just create a new doc and copy styles? No, headers are hard to copy.
    
    # Let's try this: 
    # Delete all paragraphs except the first one.
    while len(output_doc.paragraphs) > 1:
        p = output_doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
    
    # Delete all tables
    for table in output_doc.tables:
        table._element.getparent().remove(table._element)

    # Now add content from source
    first_p = True
    for p in source_doc.paragraphs:
        if not p.text.strip():
            continue
            
        if first_p:
            target_p = output_doc.paragraphs[0]
            first_p = False
        else:
            target_p = output_doc.add_paragraph()
            
        target_p.text = p.text
        
        # Simple heuristic for headings
        text = p.text.strip()
        if text.isupper() and len(text) < 100:
            copy_paragraph_style(target_p, "Heading 1", output_doc)
        elif any(text.startswith(f"{i}. ") for i in range(1, 10)) or any(text.startswith(f"{i}.{j}") for i in range(1, 10) for j in range(1, 10)):
            copy_paragraph_style(target_p, "Heading 2", output_doc)
        else:
            copy_paragraph_style(target_p, "Normal", output_doc)
            
        # Copy alignment if possible
        target_p.alignment = p.alignment

    # Copy tables
    for source_table in source_doc.tables:
        new_table = output_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
        # Copy table style if available
        try:
            new_table.style = template_doc.styles['Table Grid']
        except:
            pass
            
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text

    output_doc.save(output_path)
    print(f"Saved formatted report to {output_path}")

format_report()
