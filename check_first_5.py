import docx

doc = docx.Document("Intenship001.docx")
for i in range(min(5, len(doc.paragraphs))):
    print(f"P{i}: {repr(doc.paragraphs[i].text)}")
