import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def rgb_to_str(rgb):
    if rgb is None: return "None"
    return f"#{rgb.rgb:06X}"

def analyze_doc(path, name, max_paras=None):
    print(f"\n{'='*80}")
    print(f"ANALYZING: {name}")
    print(f"{'='*80}")
    doc = docx.Document(path)
    
    paras = doc.paragraphs if max_paras is None else doc.paragraphs[:max_paras]
    
    for i, p in enumerate(paras):
        text = p.text
        align = p.alignment
        align_name = {
            None: "None",
            0: "LEFT",
            1: "CENTER", 
            2: "RIGHT",
            3: "JUSTIFY",
        }.get(align, str(align))
        
        pf = p.paragraph_format
        sb = pf.space_before
        sa = pf.space_after
        ls = pf.line_spacing
        li = pf.left_indent
        
        print(f"\n[P{i}] Style={p.style.name} | Align={align_name} | SB={sb} SA={sa} LS={ls} LI={li}")
        print(f"  Text: {repr(text[:120])}")
        
        for j, run in enumerate(p.runs[:2]):
            fn = run.font.name
            fs = run.font.size
            fb = run.font.bold
            fi = run.font.italic
            fu = run.font.underline
            print(f"  Run{j}: Font={fn} Sz={fs} Bold={fb} Italic={fi} Underline={fu} Text={repr(run.text[:60])}")
    
    # Tables
    print(f"\n--- TABLES ({len(doc.tables)} total) ---")
    for i, table in enumerate(doc.tables[:3]):
        print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols, Style={table.style.name}")
        for ri, row in enumerate(table.rows[:3]):
            for ci, cell in enumerate(row.cells[:3]):
                print(f"  [{ri},{ci}]: {repr(cell.text[:60])}")
    
    # Key Styles
    print(f"\n--- KEY STYLES ---")
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'List Paragraph']:
        try:
            s = doc.styles[style_name]
            f = s.font
            pf = s.paragraph_format
            print(f"Style '{style_name}': Font={f.name} Size={f.size} Bold={f.bold} | SpaceBefore={pf.space_before} SpaceAfter={pf.space_after} LineSpacing={pf.line_spacing} Align={pf.alignment}")
        except Exception as e:
            print(f"Style '{style_name}': {e}")

    # Sections
    print(f"\n--- SECTIONS ---")
    for i, sec in enumerate(doc.sections):
        print(f"Section {i}: PageWidth={sec.page_width}, PageHeight={sec.page_height}")
        print(f"  Margins: L={sec.left_margin} R={sec.right_margin} T={sec.top_margin} B={sec.bottom_margin}")
        print(f"  Header paragraphs:")
        for hp in sec.header.paragraphs:
            print(f"    {repr(hp.text)} | Align={hp.alignment}")
            for run in hp.runs[:2]:
                print(f"    Run: Font={run.font.name} Sz={run.font.size} Bold={run.font.bold}")
        print(f"  Footer paragraphs:")
        for fp in sec.footer.paragraphs:
            print(f"    {repr(fp.text)} | Align={fp.alignment}")
            for run in fp.runs[:2]:
                print(f"    Run: Font={run.font.name} Sz={run.font.size} Bold={run.font.bold}")

analyze_doc("Internship_Cover_Pages_template (Final).docx", "COVER PAGES TEMPLATE")
