import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def inspect_docx(file_path):
    doc = docx.Document(file_path)
    print(f"--- Inspection of {file_path} ---")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Check headers/footers
    for section in doc.sections:
        print(f"Header: {section.header.is_linked_to_previous}")
        print(f"Footer: {section.footer.is_linked_to_previous}")

    # Check first few paragraphs for styles
    for i, p in enumerate(doc.paragraphs[:10]):
        print(f"P{i} Style: {p.style.name} | Text: {p.text[:50]}...")
        if p.runs:
            run = p.runs[0]
            print(f"  Run 0 Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}")

    # Check tables
    print(f"Number of tables: {len(doc.tables)}")

print("Inspecting Intenship001.docx")
inspect_docx("Intenship001.docx")
print("\nInspecting Landslide report.docx")
inspect_docx("Landslide report.docx")
