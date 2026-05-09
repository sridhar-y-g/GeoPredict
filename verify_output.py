import docx

doc = docx.Document("Landslide_Report_Formatted.docx")
print("First 5 paragraphs of new doc:")
for i in range(5):
    p = doc.paragraphs[i]
    print(f"[{i}] Style: {p.style.name} | Text: {repr(p.text)}")
    if p.runs:
        print(f"    Font: {p.runs[0].font.name}, Bold: {p.runs[0].font.bold}")

print("\nHeader check:")
for section in doc.sections:
    for p in section.header.paragraphs:
        print(f"Header P: {repr(p.text)}")
