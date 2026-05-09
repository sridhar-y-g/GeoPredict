import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def format_report():
    template_path = "Intenship001.docx"
    source_path = "Landslide report.docx"
    output_path = "Landslide_Report_Formatted.docx"

    template_doc = docx.Document(template_path)
    source_doc = docx.Document(source_path)
    
    # Create output doc by copying template to keep styles and headers
    output_doc = docx.Document(template_path)
    
    # Clear all paragraphs and tables from the body
    # We keep the headers/footers which are linked to the section
    for p in output_doc.paragraphs:
        p.text = ""
    while len(output_doc.paragraphs) > 1:
        p = output_doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
    
    for table in output_doc.tables:
        table._element.getparent().remove(table._element)

    # Map content
    # P0 is usually the title
    title_p = output_doc.paragraphs[0]
    title_p.text = source_doc.paragraphs[0].text
    title_p.style = output_doc.styles['Title'] if 'Title' in output_doc.styles else output_doc.styles['Heading 1']
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, p in enumerate(source_doc.paragraphs[1:]):
        if not p.text.strip():
            continue
            
        new_p = output_doc.add_paragraph()
        new_p.text = p.text
        
        text = p.text.strip()
        # Heuristic for headings
        if text.isupper() and len(text) < 100:
            new_p.style = output_doc.styles['Heading 1']
        elif any(text.startswith(f"{i}. ") for i in range(1, 15)) or any(text.startswith(f"{i}.{j}") for i in range(1, 15) for j in range(1, 15)):
            new_p.style = output_doc.styles['Heading 2']
        else:
            new_p.style = output_doc.styles['Normal']
            
        # Copy formatting from runs (bold, italic)
        # Note: This is simplified. Ideally we should iterate through runs.
        if p.runs:
            for source_run in p.runs:
                # We already added the full text to the paragraph, so we can't easily 
                # replicate runs unless we clear text and add runs one by one.
                pass

    # Better approach for runs:
    # Clear paragraph text and add runs
    for i, new_p in enumerate(output_doc.paragraphs):
        if i == 0: continue # Skip title for now or handle it
        
        # This is getting complex, let's just do it for paragraphs we add
    
    # Let's rewrite the loop to handle runs
    # Clear all paragraphs again
    for p in output_doc.paragraphs:
        p.text = ""
    while len(output_doc.paragraphs) > 1:
        p = output_doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    for i, source_p in enumerate(source_doc.paragraphs):
        if not source_p.text.strip() and i > 0:
            continue
            
        if i == 0:
            new_p = output_doc.paragraphs[0]
        else:
            new_p = output_doc.add_paragraph()
            
        text = source_p.text.strip()
        if i == 0:
            new_p.style = output_doc.styles['Heading 1']
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.isupper() and len(text) < 100:
            new_p.style = output_doc.styles['Heading 1']
        elif any(text.startswith(f"{i}. ") for i in range(1, 15)) or any(text.startswith(f"{i}.{j}") for i in range(1, 15) for j in range(1, 15)):
            new_p.style = output_doc.styles['Heading 2']
        else:
            new_p.style = output_doc.styles['Normal']
            
        for source_run in source_p.runs:
            new_run = new_p.add_run(source_run.text)
            new_run.bold = source_run.bold
            new_run.italic = source_run.italic
            new_run.underline = source_run.underline
            # Font size and name will be inherited from the style we set

    # Copy tables
    for source_table in source_doc.tables:
        new_table = output_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
        try:
            new_table.style = template_doc.styles['Table Grid']
        except:
            pass
        for r, row in enumerate(source_table.rows):
            for c, cell in enumerate(row.cells):
                new_table.cell(r, c).text = cell.text

    output_doc.save(output_path)
    print(f"Saved to {output_path}")

format_report()
