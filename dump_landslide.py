import docx

def dump_all(path, name):
    print(f"\n{'='*80}")
    print(f"ALL PARAGRAPHS OF: {name}")
    print(f"{'='*80}")
    doc = docx.Document(path)
    for i, p in enumerate(doc.paragraphs):
        print(f"[{i}] {repr(p.text)}")
    print(f"\nTotal: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    for ti, table in enumerate(doc.tables):
        print(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                print(f"  [{ri},{ci}]: {repr(cell.text)}")

dump_all("Landslide report.docx", "LANDSLIDE REPORT (SOURCE)")
