import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

source = docx.Document('Landslide report.docx')
output = docx.Document('Intenship001.docx')  # inherit styles

# Remove all existing body content
body = output.element.body
for child in list(body):
    body.remove(child)

src_paras = source.paragraphs
src_tables = source.tables

def add_p(doc, text, style, align=None, bold=None, italic=None, size_pt=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        if bold is not None: r.bold = bold
        if italic is not None: r.italic = italic
        if size_pt is not None: r.font.size = Pt(size_pt)
    return p

def add_table_copy(doc, src_table):
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            tc = t.cell(ri, ci)
            tc.text = cell.text
    return t

CHAPTER = re.compile(r'^\d+\.\s+[A-Z][A-Z\s\'\-\(\)]+$')
SECTION  = re.compile(r'^\d+\.\d+\s+.+')

LIST_PREFIXES = (
    'To ', 'Lack of', 'Delayed ', 'Limited ', 'Absence of', 'Poor ', 'Inefficient ',
    'React.js', 'Vite', 'Tailwind', 'React Leaflet', 'Framer', 'Lucide',
    'FastAPI', 'Uvicorn', 'SQLAlchemy', 'BackgroundTasks', 'smtplib', 'JWT Auth',
    'Rainfall intensity', 'Soil saturation', 'Humidity', 'Terrain class', 'Slope-based',
    'User accounts', 'Authentication data', 'OTP ', 'User profiles', 'Notification pref',
    'Secure login', 'JWT token', 'Role-based', 'Protected routes', 'Password enc',
    'MODIS', 'Dynamic image', 'Geospatial vis', 'Severe storms', 'Floods', 'Landslides',
    'Volcanic', 'Terrain slope', 'Rainfall acc', 'Geographical susc', 'Surface sat',
    'User enters', 'Frontend sends', 'Backend fetches', 'Risk engine', 'Dynamic risk',
    'Risk category', 'Results are', 'Alerts are', 'Email notif',
    'Real-time env', 'Geospatial haz', 'Dynamic risk pred', 'Integration of sat',
    'Automated em', 'Secure user auth', 'Scalable sys',
    'Sentinel-2', 'Digital Elevation', 'Slope-based terrain', 'Pixel-level',
    'terrain sus', 'environmental inf', 'geospatial analysis', 'semantic seg',
    'Dashboard vis', 'Interactive GIS', 'User authentication', 'Alert popups',
    'Satellite imagery', 'Environmental tele',
    'environmental telemetry proc', 'API integr', 'risk pred', 'authentication',
    'database oper', 'email disp', 'rainfall mon', 'humidity anal', 'weather cond',
    'Mountain geo', 'SAFE', 'WATCH', 'WARNING', 'CRITICAL',
    'user reg', 'secure log', 'OTP verif', 'protected', 'role-based',
    'visual alerts', 'emergency email', 'telemetry data', 'API response',
    'risk score gen', 'authentication testing', 'alert gen', 'dashboard func',
    'interactive geo', 'telemetry info', 'environmental anal', 'hazard vis',
    'dynamic mon', 'rainfall intens', 'humidity perc', 'weather cond', 'location-based',
    'heavy rainfall', 'high humidity', 'mountainous', 'satellite terrain', 'remote sensing data',
    'floods', 'storms', 'landslides', 'severe env', 'visual emergency',
    'HTML-based email', 'telemetry det', 'secure user reg', 'login func',
    'JWT token auth', 'role-based admin', 'user mon', 'account man', 'profile man',
    'administrative', 'frontend vis', 'backend proc', 'environmental APIs',
    'geospatial mon', 'authentication mech', 'automated alerts',
    'AI-assisted', 'geospatial intel', 'disaster man',
    'network latency', 'API response del', 'timeout', 'inconsistent',
    'coordinate', 'terrain class', 'hazard vis', 'map rend',
    'rainfall', 'humidity', 'terrain elev', 'slope inf',
    'SMTP', 'HTML email', 'background task', 'alert sync',
    'animations', 'component rend', 'map upd', 'telemetry ref',
    'dataset analysis', 'ML team', 'frontend/backend', 'documentation',
    'full-stack sys', 'geospatial tech', 'collaborative',
    'interactive geospatial vis', 'dynamic landslide', 'secure auth',
    'automated emergency', 'disaster event',
    'GIS sys', 'remote sensing', 'full-stack eng', 'environmental anal', 'intelligent dis',
    'frontend and backend', 'API-driven', 'environmental mon',
    'Convolutional', 'U-Net', 'Transformer',
    'soil moisture', 'ground vibr', 'rainfall sens', 'terrain displ',
    'provide instant', 'support offline', 'enable loc', 'improve access',
    'advanced GIS', '3D terrain', 'elevation heat', 'real-time satellite',
    'AWS', 'Microsoft Azure', 'Google Cloud', 'scalability', 'availability',
    'distributed', 'disaster resil', 'disaster management auth',
    'emergency response', 'public alert',
    'predictive trend', 'anomaly det', 'historical env', 'adaptive learn',
    'Interactive dashboard', 'GIS-based', 'Responsive user', 'Hazard analytics',
    'Alert modals', 'telemetry proc', 'prediction logic', 'NASA/OpenWeather',
    'database comm',
    'User ', 'Profile ', 'OTP\n',
    'monitor hazard', 'identify high-risk', 'generate timely', 'improve situational',
    'rainfall amount', 'slope angle', 'soil saturation',
    'baseline risk', 'environmental amplif', 'terrain inf',
    'simulated det', 'bounding boxes', 'confidence scores', 'landslide region',
    'weather severity', 'susceptibility mod',
    'Western Ghats', 'Himalayan', 'Eastern Ghats', 'normal plains',
    'storm conditions', 'precipitation intens',
    'soil instability', 'terrain saturation', 'landslide prob',
    'Interactive dashboard', 'GIS-based map', 'Hazard analytics', 'Alert modals',
    'user accounts', 'profile info', 'OTP records', 'notification pref', 'Database Tables',
    'JWT tokens', 'bcrypt', 'protected API',
    'the dashboard displays', 'emergency notifications are', 'HTML email alerts',
    'risk score ', 'terrain info', 'location details',
    'API valid', 'telemetry anal', 'risk score verif',
    'Sending pred', 'Fetching env', 'Retrieving sat', 'Managing auth', 'Receiving risk',
    'Request proc', 'Environmental data anal', 'Risk comp',
    'The frontend displays critical', 'The backend dispatches',
)

SUBHEADING = {
    'Risk Categories', 'Technologies Used', 'Backend Components', 'Main Database Tables',
    'Purpose', 'Parameters Extracted', 'Working Process', 'Features', 'Events Tracked',
    'Parameters Considered', 'Rainfall Analysis', 'Humidity and Soil Saturation',
    'Terrain Classification and Geofencing', 'Warning Thresholds', 'Frontend Features',
    'Database Tables', 'Risk Classification Categories',
    'OpenWeatherMap API', 'NASA APIs', 'NASA GIBS', 'NASA EONET', 'NASA LHASA',
    'Used for:', 'Risk Categories', 'Backend Components',
}

# Add cover page info
add_p(output, 'PREDICTIVE INTELLIGENCE FOR LANDSLIDE DISASTER PREVENTION: AN AI-POWERED EARLY WARNING SYSTEM USING HYBRID DATA SOURCES',
      'Heading 1', WD_ALIGN_PARAGRAPH.CENTER)
add_p(output, '', 'Normal')
add_p(output, 'Cohort Number: 17', 'Normal', WD_ALIGN_PARAGRAPH.CENTER, False)
add_p(output, '', 'Normal')
add_p(output, 'Team Members and Roles', 'Heading 2', WD_ALIGN_PARAGRAPH.CENTER)
add_table_copy(output, src_tables[0])
add_p(output, '', 'Normal')
add_p(output, 'Cohort Leaders', 'Heading 2', WD_ALIGN_PARAGRAPH.CENTER)
add_table_copy(output, src_tables[1])
add_p(output, '', 'Normal')

# Process main content (from paragraph index 22 = "1. INTRODUCTION")
for para in src_paras[22:]:
    text = para.text.strip()
    if not text:
        continue

    # Figure caption
    if text.startswith('Figure '):
        add_p(output, text, 'Normal', WD_ALIGN_PARAGRAPH.CENTER, False, True, 10)
        continue

    # Chapter heading
    if CHAPTER.match(text):
        add_p(output, text, 'Heading 1')
        continue

    # Section heading
    if SECTION.match(text):
        add_p(output, text, 'Heading 2')
        continue

    # Subheadings
    if text in SUBHEADING:
        p = output.add_paragraph(style='Normal')
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        continue

    # List items (short, no period at end, starts with known prefix)
    is_list = (
        len(text) < 120 and
        not text.endswith(':') and
        any(text.startswith(s) for s in LIST_PREFIXES)
    )
    if is_list:
        p = output.add_paragraph(style='List Paragraph')
        r = p.add_run(text)
        continue

    # Body text - copy runs
    p = output.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if para.runs:
        for run in para.runs:
            nr = p.add_run(run.text)
            nr.bold = run.bold
            nr.italic = run.italic
    else:
        p.add_run(text)

# Risk category tables
add_p(output, '', 'Normal')
add_table_copy(output, src_tables[2])
add_p(output, '', 'Normal')
add_table_copy(output, src_tables[3])

output.save('GeoPredict_Internship_Report.docx')
print('SUCCESS: GeoPredict_Internship_Report.docx created!')
